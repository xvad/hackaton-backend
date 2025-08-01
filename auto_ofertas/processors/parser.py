from docx import Document
import re
from typing import Dict, Any
import os
import PyPDF2
import pdfplumber

def detectar_tipo_archivo(path: str) -> str:
    """
    Detecta el tipo de archivo basado en su extensión
    """
    extension = os.path.splitext(path)[1].lower()
    if extension == '.docx':
        return 'docx'
    elif extension == '.pdf':
        return 'pdf'
    else:
        raise ValueError(f"Formato de archivo no soportado: {extension}. Solo se soportan .docx y .pdf")

def extraer_texto_pdf(path: str) -> str:
    """
    Extrae texto de un archivo PDF usando pdfplumber para mejor calidad
    """
    content = ""
    
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"
    except Exception as e:
        print(f"Error con pdfplumber: {e}")
        # Fallback a PyPDF2
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"
        except Exception as e2:
            print(f"Error con PyPDF2: {e2}")
            raise ValueError(f"No se pudo extraer texto del PDF: {path}")
    
    return content

def extraer_texto_docx(path: str) -> str:
    """
    Extrae texto de un archivo DOCX
    """
    doc = Document(path)
    content = ""
    
    # Extraer texto de párrafos con mejor manejo
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            # Preservar saltos de línea importantes
            if paragraph.style.name.startswith('Heading'):
                content += f"\n\n{text}\n"
            else:
                content += text + "\n"
    
    # Extraer texto de tablas con mejor estructura
    for table in doc.tables:
        content += "\n--- TABLA ---\n"
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_content.append(cell_text)
            if row_content:
                content += " | ".join(row_content) + "\n"
        content += "--- FIN TABLA ---\n"
    
    # Extraer texto de headers y footers si están disponibles
    try:
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        content += f"HEADER: {paragraph.text.strip()}\n"
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        content += f"FOOTER: {paragraph.text.strip()}\n"
    except:
        pass  # Algunos documentos pueden no tener headers/footers accesibles
    
    return content

def parse_licitacion_dinamica(path: str) -> Dict[str, Any]:
    """
    Parsea un documento de licitación (DOCX o PDF) y extrae su contenido estructurado
    """
    # Detectar tipo de archivo
    tipo_archivo = detectar_tipo_archivo(path)
    
    # Extraer texto según el tipo de archivo
    if tipo_archivo == 'docx':
        content = extraer_texto_docx(path)
    elif tipo_archivo == 'pdf':
        content = extraer_texto_pdf(path)
    else:
        raise ValueError(f"Tipo de archivo no soportado: {tipo_archivo}")
    
    # Detectar secciones por múltiples patrones mejorados
    secciones = {}
    seccion_actual = None
    buffer = []
    
    lines = content.split('\n')
    
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
        
        # Patrones mejorados para detectar títulos de sección
        es_titulo = False
        
        # Patrón 1: Todo en mayúsculas y más de 3 caracteres
        if line_strip.isupper() and len(line_strip) > 3:
            es_titulo = True
        
        # Patrón 2: Termina en dos puntos y tiene mayúsculas
        elif re.match(r"^[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+:$", line_strip):
            es_titulo = True
        
        # Patrón 3: Números seguidos de punto y texto (1. TÍTULO)
        elif re.match(r"^\d+\.\s*[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+$", line_strip):
            es_titulo = True
        
        # Patrón 4: Números romanos seguidos de punto (I. TÍTULO)
        elif re.match(r"^[IVX]+\.\s*[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+$", line_strip):
            es_titulo = True
        
        # Patrón 5: Letras seguidas de punto (A. TÍTULO)
        elif re.match(r"^[A-Z]\.\s*[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+$", line_strip):
            es_titulo = True
        
        # Patrón 6: Palabras clave comunes en licitaciones (expandido)
        palabras_clave = [
            "OBJETIVO", "ALCANCE", "REQUISITOS", "ESPECIFICACIONES", "PLAZOS", 
            "PRESUPUESTO", "CRITERIOS", "EVALUACIÓN", "CONDICIONES", "GARANTÍAS",
            "METODOLOGÍA", "EQUIPO", "EXPERIENCIA", "REFERENCIAS", "ENTREGABLES",
            "CRONOGRAMA", "FACTORES", "RIESGOS", "CALIDAD", "SOPORTE",
            "DESCRIPCIÓN", "CARACTERÍSTICAS", "FUNCIONALIDADES", "USUARIOS",
            "PERFILES", "ROLES", "PERMISOS", "INTEGRACIÓN", "DESARROLLO",
            "IMPLEMENTACIÓN", "CAPACITACIÓN", "DOCUMENTACIÓN", "PRUEBAS",
            "DESPLIEGUE", "MANTENIMIENTO", "SERVICIOS", "PRODUCTOS",
            "SOLUCIÓN", "SISTEMA", "PLATAFORMA", "APLICACIÓN", "SOFTWARE",
            "HARDWARE", "INFRAESTRUCTURA", "TECNOLOGÍA", "ARQUITECTURA",
            "BASE DE DATOS", "INTERFAZ", "API", "WEB", "MÓVIL", "CLOUD",
            "SEGURIDAD", "BACKUP", "RESPALDO", "MONITOREO", "REPORTES",
            "ANÁLISIS", "ESTUDIO", "DIAGNÓSTICO", "PLAN", "ESTRATEGIA",
            "PROCESO", "PROCEDIMIENTO", "POLÍTICA", "ESTÁNDAR", "NORMATIVA"
        ]
        
        if any(palabra in line_strip.upper() for palabra in palabras_clave):
            es_titulo = True
        
        # Patrón 7: Líneas que parecen títulos por su formato (corta y con mayúsculas)
        elif (len(line_strip) < 100 and 
              len(line_strip.split()) <= 8 and 
              line_strip[0].isupper() and 
              not line_strip.endswith('.') and
              not line_strip.endswith(',') and
              not re.search(r'\d+%', line_strip)):  # No porcentajes
            es_titulo = True
        
        # Patrón 8: Líneas que contienen palabras específicas de licitaciones
        palabras_licitacion = [
            "LICITACIÓN", "CONCURSO", "CONVOCATORIA", "BASES", "TÉRMINOS",
            "CONDICIONES", "REQUISITOS", "ESPECIFICACIONES", "PLIEGO",
            "PROPUESTA", "OFERTA", "PRESENTACIÓN", "EVALUACIÓN", "SELECCIÓN"
        ]
        
        if any(palabra in line_strip.upper() for palabra in palabras_licitacion):
            es_titulo = True
        
        if es_titulo:
            if seccion_actual and buffer:
                contenido_seccion = '\n'.join(buffer).strip()
                if contenido_seccion and len(contenido_seccion) > 10:
                    secciones[seccion_actual] = contenido_seccion
                buffer = []
            seccion_actual = line_strip.rstrip(':').strip()
        else:
            buffer.append(line_strip)
    
    # Agregar la última sección
    if seccion_actual and buffer:
        contenido_seccion = '\n'.join(buffer).strip()
        if contenido_seccion and len(contenido_seccion) > 10:
            secciones[seccion_actual] = contenido_seccion
    
    # Si no se detectaron secciones, intentar dividir por párrafos largos
    if not secciones:
        parrafos = [p.strip() for p in content.split('\n\n') if p.strip()]
        if len(parrafos) > 1:
            for i, parrafo in enumerate(parrafos[:10]):  # Máximo 10 secciones
                if len(parrafo) > 50:  # Solo párrafos sustanciales
                    secciones[f'Sección_{i+1}'] = parrafo
        else:
            # Si todo el contenido está en un solo bloque, dividirlo
            if len(content.strip()) > 200:
                secciones['contenido'] = content.strip()
    
    # Limpiar y mejorar secciones
    secciones_limpias = {}
    for seccion, contenido in secciones.items():
        contenido_limpio = contenido.strip()
        if contenido_limpio and len(contenido_limpio) > 20:  # Mínimo 20 caracteres
            # Limpiar líneas vacías múltiples
            contenido_limpio = re.sub(r'\n\s*\n\s*\n', '\n\n', contenido_limpio)
            secciones_limpias[seccion] = contenido_limpio
    
    # Si aún no hay secciones, crear una sección con todo el contenido
    if not secciones_limpias and content.strip():
        secciones_limpias['contenido_completo'] = content.strip()
    
    return secciones_limpias
